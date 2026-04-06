#!/usr/bin/env python3
"""
09_write_targeting_memo.py

Generates a strategic voter targeting memo as a .docx file for the
Scott Wiener for Congress campaign (CA CD-11, June 2026 Democratic primary).

Output: turnout_model/output/targeting_memo.docx
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
WIENER_BLUE = RGBColor(0x1B, 0x2A, 0x4A)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_BG = "DBE5F1"   # alternating row shading
HEADER_BG = "1B2A4A"       # table header background

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output"
OUTPUT_FILE = OUTPUT_DIR / "targeting_memo.docx"


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------
def set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, font_size=10, color=DARK_GRAY,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"
    run.font.color.rgb = color
    run.bold = bold


def add_table_with_style(doc, headers, rows, col_widths=None):
    """Create a formatted table with header row and alternating shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, header, bold=True, font_size=9, color=WHITE,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_BLUE_BG)
            bold = (c_idx == 0)  # bold the first column (labels)
            set_cell_text(cell, val, bold=bold, font_size=9,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Column widths
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    return table


def add_heading_styled(doc, text, level=1):
    """Add a heading with Wiener blue color."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = WIENER_BLUE
        run.font.name = "Calibri"
    return heading


def add_body(doc, text, bold=False, italic=False, space_after=6):
    """Add a body paragraph with standard formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.size = Pt(11)
        run_b.font.name = "Calibri"
        run_b.font.color.rgb = DARK_GRAY
        run_b.bold = True
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_GRAY
    return p


def add_sub_heading(doc, text):
    """Add a bold sub-heading line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = WIENER_BLUE
    run.bold = True
    return p


def add_confidential_watermark(doc):
    """Add a CONFIDENTIAL diagonal watermark to the document header."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Create the watermark using WordprocessingML shape in header
    watermark_xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        '     xmlns:v="urn:schemas-microsoft-com:vml"'
        '     xmlns:o="urn:schemas-microsoft-com:office:office"'
        '     xmlns:w10="urn:schemas-microsoft-com:office:word">'
        '  <w:rPr>'
        '    <w:noProof/>'
        '  </w:rPr>'
        '  <w:pict>'
        '    <v:shapetype id="_x0000_t136" coordsize="21600,21600"'
        '      o:spt="136" adj="10800"'
        '      path="m@7,l@8,m@5,21600l@6,21600e">'
        '      <v:formulas>'
        '        <v:f eqn="sum #0 0 10800"/>'
        '        <v:f eqn="prod #0 2 1"/>'
        '        <v:f eqn="sum 21600 0 @1"/>'
        '        <v:f eqn="sum 0 0 @2"/>'
        '        <v:f eqn="sum 21600 0 @3"/>'
        '        <v:f eqn="if @0 @3 0"/>'
        '        <v:f eqn="if @0 21600 @1"/>'
        '        <v:f eqn="if @0 0 @2"/>'
        '        <v:f eqn="if @0 @4 21600"/>'
        '        <v:f eqn="mid @5 @6"/>'
        '        <v:f eqn="mid @8 @5"/>'
        '        <v:f eqn="mid @7 @8"/>'
        '        <v:f eqn="mid @6 @7"/>'
        '        <v:f eqn="sum @6 0 @5"/>'
        '      </v:formulas>'
        '      <v:path textpathok="t" o:connecttype="custom"'
        '        o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800"'
        '        o:connectangles="270,180,90,0"/>'
        '      <v:textpath on="t" fitshape="t"/>'
        '      <v:handles>'
        '        <v:h position="#0,bottomRight" xrange="6629,14971"/>'
        '      </v:handles>'
        '      <o:lock v:ext="edit" text="t" shapetype="t"/>'
        '    </v:shapetype>'
        '    <v:shape id="PowerPlusWaterMarkObject"'
        '      o:spid="_x0000_s2049"'
        '      type="#_x0000_t136"'
        '      style="position:absolute;margin-left:0;margin-top:0;'
        '             width:500pt;height:100pt;rotation:315;'
        '             z-index:-251658752;mso-position-horizontal:center;'
        '             mso-position-horizontal-relative:margin;'
        '             mso-position-vertical:center;'
        '             mso-position-vertical-relative:margin"'
        '      o:allowincell="f"'
        '      fillcolor="#C0C0C0"'
        '      stroked="f">'
        '      <v:fill opacity=".25"/>'
        '      <v:textpath style="font-family:&quot;Calibri&quot;;'
        '        font-size:1pt" string="CONFIDENTIAL"/>'
        '      <w10:wrap anchorx="margin" anchory="margin"/>'
        '    </v:shape>'
        '  </w:pict>'
        '</w:r>'
    )

    # Add watermark paragraph to header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p._p.append(parse_xml(watermark_xml))


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------
def build_memo():
    """Build the complete targeting memo document."""
    doc = Document()

    # -- Page setup ----------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # -- Default font --------------------------------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Configure heading styles
    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_font = h_style.font
        h_font.name = "Calibri"
        h_font.color.rgb = WIENER_BLUE
        if level == 1:
            h_font.size = Pt(16)
        elif level == 2:
            h_font.size = Pt(14)
        else:
            h_font.size = Pt(12)

    # -- Watermark -----------------------------------------------------------
    try:
        add_confidential_watermark(doc)
    except Exception:
        pass  # Watermark is optional; don't fail the memo

    # ========================================================================
    # HEADER / TITLE PAGE
    # ========================================================================
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run("CONFIDENTIAL")
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.bold = True

    title_p2 = doc.add_paragraph()
    title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p2.paragraph_format.space_after = Pt(6)
    run2 = title_p2.add_run("CD-11 Voter Targeting Universe Memo")
    run2.font.size = Pt(24)
    run2.font.name = "Calibri"
    run2.font.color.rgb = WIENER_BLUE
    run2.bold = True

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(20)
    run3 = subtitle_p.add_run("Scott Wiener for Congress \u2014 June 2026 Primary")
    run3.font.size = Pt(16)
    run3.font.name = "Calibri"
    run3.font.color.rgb = MEDIUM_GRAY

    # Horizontal rule via bottom border on an empty paragraph
    rule_p = doc.add_paragraph()
    rule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = rule_p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="1B2A4A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Meta info
    for label, value in [
        ("Date:", "April 2026"),
        ("To:", "Campaign Leadership"),
        ("From:", "Data & Analytics Team"),
        ("RE:", "Voter Targeting Universe \u2014 CA CD-11 Democratic Primary"),
    ]:
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        meta_p.paragraph_format.space_after = Pt(2)
        meta_p.paragraph_format.left_indent = Inches(1.5)
        r1 = meta_p.add_run(label + "  ")
        r1.font.size = Pt(11)
        r1.font.name = "Calibri"
        r1.font.color.rgb = WIENER_BLUE
        r1.bold = True
        r2 = meta_p.add_run(value)
        r2.font.size = Pt(11)
        r2.font.name = "Calibri"
        r2.font.color.rgb = DARK_GRAY

    doc.add_page_break()

    # ========================================================================
    # 1. EXECUTIVE SUMMARY
    # ========================================================================
    add_heading_styled(doc, "1. Executive Summary", level=1)

    add_body(doc,
        "This memo presents the results of our voter targeting model for "
        "California's 11th Congressional District (San Francisco, Supervisor "
        "Districts 1\u201310) ahead of the June 2026 Democratic primary election. "
        "The model assigns every registered voter a composite support score "
        "(0\u2013100) calibrated against EMC Research polling (February 2026 "
        "and September 2025) and segments the electorate into three strategic "
        "universes: Base, Persuasion, and Opposition."
    )

    add_body(doc, "Key Findings:", bold=True, space_after=4)

    add_bullet(doc, "448,376 registered voters in CD-11 (Supervisor Districts 1\u201310)",
               bold_prefix="Total Electorate: ")
    add_bullet(doc, "Base 177,212 (39.5%) | Persuasion 168,183 (37.5%) | Opposition 102,981 (23.0%)",
               bold_prefix="Universe Breakdown: ")
    add_bullet(doc, "Saikat Chakrabarti \u2014 progressive challenger with deep-left coalition",
               bold_prefix="Primary Opponent: ")
    add_bullet(doc,
        "Scott\u2019s base is concentrated among center-left, college-educated, "
        "LGBTQ+, White/API voters in Supervisor Districts 2, 3, 6, 7, and 8.",
        bold_prefix="Base Profile: ")
    add_bullet(doc,
        "Mobilize Base drop-off voters (66,155 supporters with low turnout "
        "probability) and persuade swing voters in Districts 1, 4, 5, and 9.",
        bold_prefix="Strategic Imperative: ")

    # Summary table
    doc.add_paragraph()
    summary_headers = ["Universe", "Voters", "Share", "Avg Score", "Avg Turnout"]
    summary_rows = [
        ["Base", "177,212", "39.5%", "76.8", "51.7%"],
        ["Persuasion", "168,183", "37.5%", "52.4", "32.0%"],
        ["Opposition", "102,981", "23.0%", "42.1", "30.3%"],
        ["TOTAL", "448,376", "100.0%", "\u2014", "\u2014"],
    ]
    add_table_with_style(doc, summary_headers, summary_rows,
                         col_widths=[1.5, 1.2, 1.0, 1.1, 1.1])

    doc.add_page_break()

    # ========================================================================
    # 2. METHODOLOGY
    # ========================================================================
    add_heading_styled(doc, "2. Methodology", level=1)

    add_body(doc,
        "Each registered voter in CD-11 receives a composite support score "
        "(0\u2013100) derived from eight weighted factors, calibrated against "
        "EMC Research polling crosstabs (February 2026, n=800; September "
        "2025, n=500, weighted 60/40). The model integrates TargetSmart "
        "modeled data with geographic, demographic, and behavioral "
        "indicators to estimate each voter\u2019s likelihood of supporting "
        "Scott Wiener in a head-to-head primary."
    )

    add_sub_heading(doc, "Scoring Factors")

    factor_headers = ["Factor", "Max Points", "Data Source"]
    factor_rows = [
        ["Ideology", "30", "TargetSmart Partisan Score + modeled ideology"],
        ["LGBTQ+ Proxy", "20", "Marriage Equality support + HRC donation flag"],
        ["Geography", "15", "Supervisor District performance benchmarks"],
        ["Race/Ethnicity", "10", "TargetSmart modeled ethnicity"],
        ["Age", "8", "Voter file date of birth"],
        ["Education", "7", "TargetSmart modeled education level"],
        ["Vote Frequency", "5", "Historical primary turnout (2018\u20132024)"],
        ["Party Registration", "5", "Voter file party affiliation"],
    ]
    add_table_with_style(doc, factor_headers, factor_rows,
                         col_widths=[1.5, 1.0, 3.5])

    doc.add_paragraph()

    add_sub_heading(doc, "Universe Thresholds")
    add_bullet(doc, "Score >= 65", bold_prefix="Base: ")
    add_bullet(doc, "Score 40\u201364", bold_prefix="Persuasion: ")
    add_bullet(doc, "Score < 40", bold_prefix="Opposition: ")

    add_body(doc,
        "Opposition voters are further segmented using TargetSmart Trump "
        "Support scores (MAGA flag) and very-progressive ideology markers "
        "(Saikat-likely). VAN voter IDs (985 real contacts) were used as "
        "soft validation; however, these IDs do not discriminate well in "
        "all-Democratic San Francisco primaries and are not incorporated "
        "directly into the scoring model.",
        italic=True
    )

    add_body(doc,
        "Note: The model relies on TargetSmart modeled scores for ideology, "
        "education, homeownership, and LGBTQ+ proxies. These are probabilistic "
        "estimates and should be interpreted accordingly.",
        italic=True
    )

    doc.add_page_break()

    # ========================================================================
    # 3. UNIVERSE PROFILES
    # ========================================================================
    add_heading_styled(doc, "3. Universe Profiles", level=1)

    # --- Base Universe ---
    add_sub_heading(doc, "Base Universe (177,212 \u2014 39.5%)")

    add_body(doc,
        "The Base universe comprises voters most likely to support Scott Wiener "
        "based on ideological alignment, demographic profile, and geographic "
        "concentration. These voters form the campaign\u2019s core coalition."
    )

    base_headers = ["Metric", "Value"]
    base_rows = [
        ["Average Support Score", "76.8"],
        ["Average Turnout Probability", "51.7%"],
        ["Democrat Registration", "77.7%"],
        ["White", "63.0%"],
        ["Asian/Pacific Islander", "29.4%"],
        ["Female", "48.9%"],
        ["Age 50+", "55.3%"],
        ["Top District: D8 (Castro/Noe Valley)", "55.3% Base"],
        ["District 7 (West Side)", "51.8% Base"],
        ["District 2 (Marina/Pac Heights)", "50.0% Base"],
    ]
    add_table_with_style(doc, base_headers, base_rows,
                         col_widths=[3.5, 2.0])

    doc.add_paragraph()
    add_body(doc,
        "Drop-off Base: 66,155 voters (37.3% of Base) are supporters with "
        "low turnout probability. These individuals are the campaign\u2019s "
        "highest-ROI mobilization targets \u2014 they support Scott but may not "
        "vote in a June primary without direct contact.",
        bold=False
    )

    doc.add_paragraph()

    # --- Persuasion Universe ---
    add_sub_heading(doc, "Persuasion Universe (168,183 \u2014 37.5%)")

    add_body(doc,
        "Persuasion voters are ideologically and demographically mixed. They "
        "lack strong signals for or against Scott and represent the campaign\u2019s "
        "opportunity to expand the coalition through targeted messaging."
    )

    pers_headers = ["Metric", "Value"]
    pers_rows = [
        ["Average Support Score", "52.4"],
        ["Average Turnout Probability", "32.0%"],
        ["Democrat Registration", "64.9%"],
        ["White", "54.0%"],
        ["Asian/Pacific Islander", "30.3%"],
        ["Female", "46.2%"],
    ]
    add_table_with_style(doc, pers_headers, pers_rows,
                         col_widths=[3.5, 2.0])

    doc.add_paragraph()
    add_sub_heading(doc, "Priority Persuasion Sub-Segments")

    seg_headers = ["Segment", "Count", "Key Message Frame"]
    seg_rows = [
        ["API Voters", "50,680", "Housing affordability + ICE enforcement protection"],
        ["Male 18\u201349", "51,081", "Anti-Trump record + housing/economic opportunity"],
        ["AD 19 Voters", "75,117", "Targeted field outreach in underperforming AD"],
        ["Strong Renters", "90,206", "Housing affordability, tenant protection record"],
        ["Jewish Community", "9,207", "Israel policy, community safety messaging"],
    ]
    add_table_with_style(doc, seg_headers, seg_rows,
                         col_widths=[1.5, 1.0, 3.5])

    doc.add_paragraph()

    # --- Opposition Universe ---
    add_sub_heading(doc, "Opposition Universe (102,981 \u2014 23.0%)")

    add_body(doc,
        "The Opposition universe includes voters unlikely to support Scott "
        "due to ideological misalignment (far-right or far-left), low "
        "engagement, or affinity for Saikat Chakrabarti. These voters are "
        "deprioritized in resource allocation."
    )

    opp_headers = ["Sub-Segment", "Count", "% of Total", "Profile"]
    opp_rows = [
        ["MAGA / Trump-aligned", "37,232", "8.3%", "High Trump Support score; unlikely Dem primary voters"],
        ["Saikat-likely", "9,331", "2.1%", "Very progressive; aligned with Chakrabarti coalition"],
        ["Conservative", "3,994", "0.9%", "Non-MAGA right; low Dem primary engagement"],
        ["Low-score Residual", "52,424", "11.7%", "Weak signals across all factors; low engagement"],
    ]
    add_table_with_style(doc, opp_headers, opp_rows,
                         col_widths=[1.6, 0.9, 0.9, 2.8])

    doc.add_paragraph()
    opp_demo_text = (
        "Opposition demographics: 38.9% Democrat, 26.7% Republican, "
        "44.6% White, 41.9% age 18\u201334. The relatively high share of "
        "young voters reflects low-engagement registrants who score poorly "
        "across multiple factors."
    )
    add_body(doc, opp_demo_text, italic=True)

    doc.add_page_break()

    # ========================================================================
    # 4. GEOGRAPHIC ANALYSIS
    # ========================================================================
    add_heading_styled(doc, "4. Geographic Analysis", level=1)

    add_body(doc,
        "The table below presents the targeting breakdown by San Francisco "
        "Supervisor District. Districts are ordered by Base share (descending) "
        "to highlight the campaign\u2019s geographic strengths and vulnerabilities."
    )

    geo_headers = [
        "District", "Description", "Total\nVoters", "Base\nCount",
        "Base %", "Persuasion %", "Opposition %"
    ]
    geo_rows = [
        ["D8",  "Castro / Noe Valley",       "58,030", "32,119", "55.3%", "33.8%", "10.9%"],
        ["D7",  "West Side / Twin Peaks",     "50,330", "26,047", "51.8%", "30.1%", "18.1%"],
        ["D2",  "Marina / Pacific Heights",   "48,283", "24,128", "50.0%", "34.7%", "15.4%"],
        ["D3",  "North Beach / Chinatown",    "42,689", "20,163", "47.2%", "39.1%", "13.6%"],
        ["D6",  "SoMa / Tenderloin",          "36,782", "16,226", "44.1%", "41.6%", "14.3%"],
        ["D9",  "Mission / Bernal Heights",   "39,395", "13,584", "34.5%", "39.0%", "26.5%"],
        ["D5",  "Haight / Western Addition",  "44,384", "13,563", "30.6%", "39.9%", "29.5%"],
        ["D1",  "Richmond",                   "46,755", "12,467", "26.7%", "42.0%", "31.3%"],
        ["D4",  "Sunset / Parkside",          "48,802", "12,648", "25.9%", "42.1%", "32.0%"],
        ["D10", "Bayview / Potrero Hill",     "32,926", "6,267",  "19.0%", "34.6%", "46.4%"],
    ]
    add_table_with_style(doc, geo_headers, geo_rows,
                         col_widths=[0.5, 1.7, 0.8, 0.8, 0.7, 0.9, 0.9])

    doc.add_paragraph()

    add_sub_heading(doc, "Key Geographic Takeaways")
    add_bullet(doc,
        "District 8 (Castro/Noe Valley) is Scott\u2019s strongest district at "
        "55.3% Base \u2014 LGBTQ+ concentration, long-standing constituent relationships.",
        bold_prefix="Stronghold: ")
    add_bullet(doc,
        "Districts 7 and 2 round out the top tier with roughly half of all "
        "voters in the Base universe.",
        bold_prefix="Core Base: ")
    add_bullet(doc,
        "District 5 (Haight/Western Addition) is the progressive stronghold "
        "most favorable to Saikat Chakrabarti. At 30.6% Base with 29.5% "
        "Opposition \u2014 the blended polling data shows stronger-than-expected "
        "support here; targeted outreach to moderate progressives is warranted.",
        bold_prefix="Competitive Progressive: ")
    add_bullet(doc,
        "District 4 (Sunset) is heavily API with the highest Persuasion share "
        "(42.1%). Culturally tailored messaging could yield significant gains.",
        bold_prefix="Persuasion Opportunity: ")
    add_bullet(doc,
        "District 10 (Bayview/Potrero Hill) has the lowest Base share (19.0%) "
        "and highest Opposition (46.4%). Resources should be allocated elsewhere.",
        bold_prefix="Deprioritize: ")

    doc.add_page_break()

    # ========================================================================
    # 5. STRATEGIC RECOMMENDATIONS
    # ========================================================================
    add_heading_styled(doc, "5. Strategic Recommendations", level=1)

    # --- A. Mobilization ---
    add_sub_heading(doc, "A. Mobilization Priority: Drop-off Base (66,155 voters)")

    add_body(doc,
        "The single highest-ROI investment this campaign can make is mobilizing "
        "the 66,155 Base voters with low turnout probability. These individuals "
        "already support Scott but are unlikely to vote in a June primary "
        "without direct outreach. Every vote earned here comes from an existing "
        "supporter \u2014 no persuasion required."
    )

    add_bullet(doc,
        "Concentrate GOTV resources (door knocks, phone calls, texts) on "
        "drop-off Base voters.",
        bold_prefix="Tactic: ")
    add_bullet(doc,
        "Districts 2, 3, 6, 7, 8 where Base concentration is highest.",
        bold_prefix="Geographic Focus: ")
    add_bullet(doc,
        "Emphasize urgency and primary-specific turnout. Frame: "
        "\u201cYour vote matters in June \u2014 this primary decides who represents SF.\u201d",
        bold_prefix="Messaging: ")
    add_bullet(doc,
        "VBM chase program targeting drop-off Base voters who requested "
        "ballots but have not returned them.",
        bold_prefix="Vote-by-Mail Chase: ")

    doc.add_paragraph()

    # --- B. Persuasion ---
    add_sub_heading(doc, "B. Persuasion Targets: Priority Segments")

    add_body(doc,
        "The Persuasion universe is large (168,183) but heterogeneous. "
        "The campaign should focus paid media, direct mail, and field "
        "resources on five high-impact sub-segments with tailored messaging."
    )

    add_bullet(doc,
        " (50,680 voters) \u2014 Housing affordability and ICE enforcement "
        "protection messaging. Culturally competent outreach in Cantonese, "
        "Mandarin, and Tagalog. Concentrate in Districts 1, 3, 4.",
        bold_prefix="1. API Voters in Persuasion")
    add_bullet(doc,
        " (51,081 voters) \u2014 Anti-Trump record, housing and economic "
        "opportunity messaging. Digital-first outreach strategy.",
        bold_prefix="2. Male 18\u201349 in Persuasion")
    add_bullet(doc,
        " (75,117 voters) \u2014 Targeted field outreach in the "
        "underperforming assembly district. Coordinate with local "
        "AD 19 Democratic clubs.",
        bold_prefix="3. AD 19 Voters")
    add_bullet(doc,
        " (90,206 voters) \u2014 Housing affordability platform, "
        "tenant protection legislative record. Emphasize Scott\u2019s "
        "authored housing bills.",
        bold_prefix="4. Strong Renters")
    add_bullet(doc,
        " (9,207 voters) \u2014 Israel policy, community safety, "
        "antisemitism response. Synagogue and community center outreach.",
        bold_prefix="5. Jewish Community")

    doc.add_paragraph()

    # --- C. Opposition Containment ---
    add_sub_heading(doc, "C. Opposition Containment")

    add_body(doc,
        "The Opposition universe should receive minimal campaign resources. "
        "Engaging these voters risks wasting limited budget and potentially "
        "motivating opposition turnout."
    )

    contain_headers = ["Segment", "Count", "Recommendation"]
    contain_rows = [
        ["Saikat-likely", "9,331", "Do not engage \u2014 resource waste; firm opposition"],
        ["MAGA Voters", "37,232", "Do not engage in primary; many won\u2019t vote in Dem primary"],
        ["Conservative", "3,994", "Minimal \u2014 ignore"],
        ["Low-score Residual", "52,424", "Some recoverable with strong messaging; low priority"],
    ]
    add_table_with_style(doc, contain_headers, contain_rows,
                         col_widths=[1.5, 1.0, 3.5])

    doc.add_paragraph()

    # --- D. Geographic Strategy ---
    add_sub_heading(doc, "D. Geographic Strategy")

    geo_strat_headers = ["Investment Level", "Districts", "Rationale"]
    geo_strat_rows = [
        ["Invest Heavily", "D7, D8", "Highest Base concentration; mobilization focus"],
        ["Invest Moderately", "D2, D3, D6", "Strong Base; ensure turnout"],
        ["Persuade", "D1, D4, D9", "Swing districts; targeted messaging opportunity"],
        ["Contest", "D5", "Progressive stronghold; compete for moderate progressives"],
        ["Deprioritize", "D10", "High opposition, low ROI"],
    ]
    add_table_with_style(doc, geo_strat_headers, geo_strat_rows,
                         col_widths=[1.3, 1.0, 3.7])

    doc.add_page_break()

    # ========================================================================
    # 6. KEY RISKS & CAVEATS
    # ========================================================================
    add_heading_styled(doc, "6. Key Risks & Caveats", level=1)

    add_body(doc,
        "The targeting model provides a strong strategic foundation, but "
        "several limitations and risks should inform how the campaign "
        "interprets and acts on these results."
    )

    add_bullet(doc,
        "The model relies on TargetSmart modeled scores (ideology, education, "
        "homeownership, LGBTQ+ proxy) which are probabilistic estimates, not "
        "actual voter preferences. Individual-level accuracy is limited.",
        bold_prefix="Modeled Data Limitations: ")

    add_bullet(doc,
        "LGBTQ+ identity is proxied through Marriage Equality support indicators "
        "and HRC donation flags. This method likely undercounts younger LGBTQ+ "
        "voters who have not donated or been surveyed.",
        bold_prefix="LGBTQ+ Proxy Imperfection: ")

    add_bullet(doc,
        "The EMC Research poll used for calibration was conducted in February "
        "2026. Voter sentiment may shift as the June primary approaches, "
        "particularly in response to endorsements, debates, and campaign events.",
        bold_prefix="Polling Shelf Life: ")

    add_bullet(doc,
        "Saikat Chakrabarti\u2019s progressive coalition could grow significantly "
        "with key endorsements (e.g., from outgoing progressive supervisors, "
        "national progressive figures). The campaign should monitor and "
        "reassess the Saikat-likely segment monthly.",
        bold_prefix="Opponent Coalition Growth: ")

    add_bullet(doc,
        "June primaries historically produce low turnout. If turnout drops "
        "below 35%, the Persuasion universe becomes less reliable \u2014 only "
        "high-propensity voters will cast ballots, compressing the effective "
        "electorate toward the Base and Opposition universes.",
        bold_prefix="Low Turnout Risk: ")

    add_bullet(doc,
        "VAN voter IDs (985 contacts) are too few and too concentrated to "
        "serve as robust model validation. The campaign should prioritize "
        "expanding voter ID efforts, particularly in Persuasion-heavy "
        "districts (D1, D4, D9).",
        bold_prefix="Limited Voter ID Validation: ")

    doc.add_paragraph()
    doc.add_paragraph()

    # Closing line
    closing_p = doc.add_paragraph()
    closing_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing_p.paragraph_format.space_before = Pt(20)

    # Horizontal rule
    rule_p2 = doc.add_paragraph()
    rule_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = rule_p2._p.get_or_add_pPr()
    pBdr2 = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="1B2A4A"/>'
        f'</w:pBdr>'
    )
    pPr2.append(pBdr2)

    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_p.paragraph_format.space_before = Pt(10)
    run_end = end_p.add_run(
        "Prepared by the Data & Analytics Team \u2014 April 2026\n"
        "Scott Wiener for Congress\n"
        "CONFIDENTIAL \u2014 Do Not Distribute"
    )
    run_end.font.size = Pt(10)
    run_end.font.name = "Calibri"
    run_end.font.color.rgb = MEDIUM_GRAY
    run_end.italic = True

    # -- Footer with page numbers -------------------------------------------
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(6)

    run_pre = fp.add_run("CONFIDENTIAL  \u2014  Scott Wiener for Congress  \u2014  Page ")
    run_pre.font.size = Pt(8)
    run_pre.font.name = "Calibri"
    run_pre.font.color.rgb = MEDIUM_GRAY

    # Page number field
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld_char_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

    run_pg = fp.add_run()
    run_pg.font.size = Pt(8)
    run_pg.font.name = "Calibri"
    run_pg.font.color.rgb = MEDIUM_GRAY
    run_pg._r.append(fld_char_begin)
    run_pg2 = fp.add_run()
    run_pg2._r.append(instr_text)
    run_pg3 = fp.add_run()
    run_pg3._r.append(fld_char_sep)
    run_pg4 = fp.add_run()
    run_pg4._r.append(fld_char_end)

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    print("Building targeting memo...")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = build_memo()
    doc.save(str(OUTPUT_FILE))

    print(f"Memo saved to: {OUTPUT_FILE}")
    print(f"File size: {OUTPUT_FILE.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
